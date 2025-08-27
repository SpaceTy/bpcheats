import os
import time
from util import read_docx_file
from llm import ask_llm
from terminal_emulator import run_command_and_capture_output
from docx import Document
from docx.shared import Inches


def main():
    """
    Main function that orchestrates the entire process:
    1. Clear the field directory
    2. Read the DOCX file
    3. Request LLM to generate commands
    4. Execute commands and generate output images
    5. Create a new DOCX with images and descriptions
    """
    print("Starting the automated process...")

    # Step 1: Clear the field directory
    print("Step 1: Clearing the field directory...")
    clear_field_directory()

    # Step 2: Read the DOCX file
    print("Step 2: Reading the DOCX file...")
    try:
        docx_content = read_docx_file("exprement number 3.docx")
        print("DOCX file read successfully.")
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return

    # Step 2: Request LLM to generate commands
    print("Step 2: Requesting LLM to generate commands...")
    system_prompt = """
    You are a Linux terminal expert. Your task is to convert the following instructions into a series of Linux commands that will accomplish the described tasks.

    Please provide only the commands, one per line, in the exact order they should be executed. Do not include any explanations, comments, or markdown formatting.

    Example format:
    mkdir lab3work
    cd lab3work
    touch aryan1 aryan2 aryan3
    echo "Line 1" > aryan1
    echo "Line 2" >> aryan1
    # ... etc

    The commands will be executed in order in a Linux terminal.
    """

    user_message = f"Convert these instructions into Linux commands:\n\n{docx_content}"

    try:
        commands_response = ask_llm(system_prompt, user_message)
        commands = [cmd.strip() for cmd in commands_response.strip().split('\n') if cmd.strip()]
        print(f"Generated {len(commands)} commands.")
        print("Commands to execute:")
        for i, cmd in enumerate(commands, 1):
            print(f"  {i}. {cmd}")
    except Exception as e:
        print(f"Error generating commands with LLM: {e}")
        return

    # Step 3: Execute commands and generate output images
    print("Step 3: Executing commands and generating output images...")
    image_paths = []
    command_descriptions = []

    # Clear the output directory
    if os.path.exists("output"):
        for file in os.listdir("output"):
            os.remove(os.path.join("output", file))

    for i, command in enumerate(commands, 1):
        print(f"Executing command {i}/{len(commands)}: {command}")
        try:
            # Add a small delay to ensure commands don't overlap
            time.sleep(1)

            # Execute command and generate output image
            image_path = run_command_and_capture_output(command)
            image_paths.append(image_path)
            command_descriptions.append(command)
            print(f"Output saved to: {image_path}")
        except Exception as e:
            print(f"Error executing command '{command}': {e}")
            # Continue with the next command

    # Step 4: Create a new DOCX with images and descriptions
    print("Step 4: Creating DOCX report with images and descriptions...")

    try:
        create_docx_report(docx_content, command_descriptions, image_paths)
        print("Final DOCX report created successfully.")
    except Exception as e:
        print(f"Error creating final report: {e}")

    print("Process completed!")


def create_docx_report(original_instructions, commands, image_paths):
    """
    Create a DOCX report with commands and their output images.

    Args:
        original_instructions (str): Original instructions from DOCX file
        commands (list): List of executed commands
        image_paths (list): List of paths to output images
    """
    # Ensure output directory exists
    output_dir = "output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created output directory: {output_dir}")

    # Create a new document
    doc = Document()

    # Add title
    doc.add_heading('Linux Command Execution Report', 0)

    # Add original instructions
    doc.add_heading('Original Instructions', 1)
    doc.add_paragraph(original_instructions)

    # Add command execution section
    doc.add_heading('Command Execution Results', 1)

    # Add each command with its image
    for i, (command, image_path) in enumerate(zip(commands, image_paths)):
        # Add command heading
        doc.add_heading(f'Command {i+1}: {command}', 2)

        # Add command description
        doc.add_paragraph(generate_command_description(command))

        # Add image if it exists
        if os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f'[Error loading image: {str(e)}]')
        else:
            doc.add_paragraph('[Image not found]')

        doc.add_paragraph()  # Add spacing

    # Save the document
    output_path = os.path.join(output_dir, "command_execution_report.docx")
    try:
        doc.save(output_path)
        print(f"DOCX report saved to: {output_path}")
    except Exception as e:
        print(f"Error saving DOCX report: {e}")
        # Try saving to current directory as fallback
        fallback_path = "command_execution_report.docx"
        try:
            doc.save(fallback_path)
            print(f"DOCX report saved to fallback location: {fallback_path}")
        except Exception as e2:
            print(f"Error saving DOCX report to fallback location: {e2}")


def generate_command_description(command):
    """
    Generate a human-readable description of what a command does by asking an LLM.

    Args:
        command (str): The command to describe

    Returns:
        str: A description of what the command does
    """
    command = command.strip()

    # Ask the LLM to describe what the command does
    system_prompt = """
    You are a Linux terminal expert. Your task is to provide a brief, human-readable description of what a given Linux command does.

    The description should be concise and clear, suitable for inclusion in a technical report.

    Example format:
    "This command creates a new directory named 'example'."
    "This command lists the contents of the current directory."
    "This command copies 'file1.txt' to 'file2.txt'."

    Do not include any explanations, comments, or markdown formatting. Just provide the description.
    """

    user_message = f"Describe what this Linux command does: {command}"

    try:
        description = ask_llm(system_prompt, user_message)
        # Clean up the response in case there's extra whitespace or newlines
        if description:
            description = description.strip()
            if description:
                return description

        return "This command was executed in the field directory to perform the required task."
    except Exception:
        # If there's any error in asking the LLM, fall back to a generic description
        return "This command was executed in the field directory to perform the required task."




def clear_field_directory():
    """
    Clear the contents of the field directory.
    """
    field_dir = "field"

    if not os.path.exists(field_dir):
        print(f"Field directory '{field_dir}' does not exist. Creating it...")
        os.makedirs(field_dir)
        return

    try:
        for item in os.listdir(field_dir):
            item_path = os.path.join(field_dir, item)
            if os.path.isfile(item_path) or os.path.islink(item_path):
                os.unlink(item_path)
            elif os.path.isdir(item_path):
                import shutil
                shutil.rmtree(item_path)
        print(f"Field directory '{field_dir}' cleared successfully.")
    except Exception as e:
        print(f"Warning: Could not clear field directory: {e}")


if __name__ == "__main__":
    main()
